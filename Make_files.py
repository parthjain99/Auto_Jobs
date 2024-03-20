import os
import sys
import shutil
from google_smart import get_match_score, get_eval, input_pdf_setup, cover_letter
from docx import Document
from docx.shared import Pt
from CSV_saving import save_to_file
from docx.shared import RGBColor

def copys(root_src_dir, root_dst_dir):
    for src_dir, dirs, files in os.walk(root_src_dir):
        dst_dir = src_dir.replace(root_src_dir, root_dst_dir, 1)
        if not os.path.exists(dst_dir):
            os.makedirs(dst_dir)
        for file_ in files:
            src_file = os.path.join(src_dir, file_)
            dst_file = os.path.join(dst_dir, file_)
            if os.path.exists(dst_file):
                os.remove(dst_file)
            shutil.copy(src_file, dst_dir)


parent_dir = "/Users/parthjain/Desktop/jobs/newS"
path = os.path.join(parent_dir, sys.argv[2])
if not os.path.exists(path):
    os.mkdir(path)

if sys.argv[1] == "BE":
    copys("/Users/parthjain/Desktop/jobs/templates/backend", path)
elif sys.argv[1] == "DS":
    copys("/Users/parthjain/Desktop/jobs/templates/Dataanalysis", path)
elif sys.argv[1] == "ML":
    copys("/Users/parthjain/Desktop/jobs/templates/machine leanring", path)

print("_________________file______")
print(path)
path_p = os.path.join(parent_dir, sys.argv[2], "Parth_Jain_Resume.pages")
path_r = os.path.join(parent_dir, sys.argv[2], "Parth_Jain_Resume.pdf")
path_c = os.path.join(parent_dir, sys.argv[2], "Parth Jain_cover_letter.docx")
print("____________Resume File________________")
print(path_p)

print("__________________JD____________________")
jd = []
while True:
    line = input("Paste Job description\n")
    if line == "quit":
        break
    jd.append(line)
jd = "\n".join(jd)
resume = input_pdf_setup(path_r)
print("_______________match Score________________")
get_match_score(jd, resume)
print("______________suggetions_________________")
get_eval(jd, resume)

doc = Document(path_c)
resume = input_pdf_setup(path_r)
doc.tables[0].rows[1].cells[1].text = cover_letter(jd, resume, sys.argv[2])
style = doc.styles.add_style('NewStyle', 1)
style.font.name = 'Calibri'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0,0,0)
# Apply the new style to each paragraph
for paragraph in doc.tables[0].rows[1].cells[1].paragraphs:
    paragraph.style = doc.styles['NewStyle']
doc.save(f'/Users/parthjain/Desktop/Jobs/news/{sys.argv[2]}/Parth Jain_cover_letter1.docx') 
save_to_file(company_name=sys.argv[1], jd=jd, role=sys.argv[2])
print("Cover letter Saved")
