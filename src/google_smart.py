from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain.chains import LLMChain, SequentialChain
from docx import Document
from pypdf import PdfReader, PdfWriter
from docx.shared import Pt
from docx.shared import RGBColor
import sys,io
sys.path.append("..")
from config.config import (challenge_prompt, intro_prompt, body_prompt, 
                           conclusion_prompt, get_match_prompt, eval_prompt,
                           FName, LName, unified_prompt)

model = ChatGoogleGenerativeAI(model="gemini-1.5-flash")


def input_pdf_setup(uploaded_file):
    if uploaded_file is not None:
        ## Convert the PDF to image
        # creating a pdf file object

        # creating a pdf reader object
        pdf_stream = io.BytesIO(uploaded_file)

        pdfReader = PdfReader(pdf_stream)

        # creating a page object
        pageObj = pdfReader.pages[0]
        txt = pageObj.extract_text()
        #extracting text from page
        return txt


def get_eval(jd, resume):
    if resume is not None:
        # pdf_content = input_pdf_setup(uploaded_file)
        eval_prompt_template = PromptTemplate.from_template(template = eval_prompt)
        eval_chain = LLMChain(prompt = eval_prompt_template ,llm = model, output_key = "eval_score")
        response = eval_chain.invoke({"resume":resume, "jd":jd})
        return response['eval_score']



def get_match_score(jd, resume):

    # pdf_content = input_pdf_setup(uploaded_file)
    get_match_prompt_template = PromptTemplate.from_template(template = get_match_prompt)
    get_match_chain = LLMChain(prompt = get_match_prompt_template ,llm = model, output_key = "match_score")
    response = get_match_chain.invoke({"resume":resume, "jd":jd})
    return response['match_score']

def cover_letter(jd, resume, company):
    challenge_prompt_template = PromptTemplate.from_template(template = challenge_prompt)
    challenge_chain = LLMChain(llm = model, prompt = challenge_prompt_template, 
                                output_key = "challenges")

    intro_prompt_template = PromptTemplate.from_template(template = intro_prompt)
    intro_chain = LLMChain(llm = model, prompt = intro_prompt_template, 
                            output_key = "hook1")

    body_prompt_template = PromptTemplate.from_template(template = body_prompt)
    body_chain = LLMChain(llm = model, prompt = body_prompt_template, 
                            output_key = "hook2")
    
    conclusion_prompt_template = PromptTemplate.from_template(template = conclusion_prompt)
    conclusion_chain = LLMChain(llm = model, prompt = conclusion_prompt_template, 
                            output_key = "hook3")
    
    final_chain  = SequentialChain(chains = [challenge_chain, intro_chain, body_chain, conclusion_chain],
                                    input_variables=["jd", "resume", "company"],
                                    output_variables=["hook1", "hook2", "hook3"])
    final = final_chain.invoke({"jd":jd, "resume": resume, "company": company})
    hook1, hook2, hook3  = final["hook1"], final["hook2"], final["hook3"]
    cv = f"""Hello Hiring team at {company},\n\n {hook1} \n\n {hook2} \n\n {hook3} \n\n Thank you for the opportunity, \n {FName} {LName}"""
    doc = Document()
    style = doc.styles.add_style('NewStyle', 1)
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.line_spacing = 0.5
    
    # Add the cover letter to the document
    for line in cv.split("\n"):
        paragraph = doc.add_paragraph(line, style = "No Spacing")
        paragraph.style = doc.styles['NewStyle']

    # Save the document to a BytesIO stream
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)  # Reset the pointer to the beginning of the stream

    # Return the byte stream content
    return cv, byte_stream.getvalue()


def cover_letter_unified_prompt(jd, resume, company):
    unified_prompt_template = PromptTemplate.from_template(template = unified_prompt)

    unified_chain = LLMChain(llm=model, prompt=unified_prompt_template)

    final = unified_chain.invoke({"jd":jd, "resume": resume, "company": company})
    cv = f"""{final['text']}"""
    doc = Document()
    style = doc.styles.add_style('NewStyle', 1)
    style.font.name = 'Calibri'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.line_spacing = 0.5
    
    # Add the cover letter to the document
    for line in cv.split("\n"):
        paragraph = doc.add_paragraph(line, style = "No Spacing")
        paragraph.style = doc.styles['NewStyle']

    # Save the document to a BytesIO stream
    byte_stream = io.BytesIO()
    doc.save(byte_stream)
    byte_stream.seek(0)  # Reset the pointer to the beginning of the stream

    # Return the byte stream content
    return cv, byte_stream.getvalue()